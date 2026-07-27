"""Reading text, tables, and figures out of whatever file the user uploaded.

One dispatcher per capability keeps ingestion independent of file format: it asks for
text, tables, and figures, and does not care whether they came from a PDF, a Word
document, a plain text file, or a bare image.
"""

import os
import shutil
import sys

sys.stdout.reconfigure(encoding="utf-8")

TEXT_SUFFIXES = {".txt", ".md", ".markdown", ".rst"}
IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp"}
SUPPORTED_SUFFIXES = {".pdf", ".docx"} | TEXT_SUFFIXES | IMAGE_SUFFIXES


class UnsupportedFileError(ValueError):
    """Raised for a file type the pipeline cannot read."""


def suffix(path: str) -> str:
    return os.path.splitext(path)[1].lower()


def is_image(path: str) -> bool:
    return suffix(path) in IMAGE_SUFFIXES


def check_supported(path: str) -> None:
    if suffix(path) not in SUPPORTED_SUFFIXES:
        supported = ", ".join(sorted(SUPPORTED_SUFFIXES))
        raise UnsupportedFileError(f"Cannot read {suffix(path) or 'this file'}. Supported: {supported}")


def load_text(path: str) -> str:
    """Extract the document's readable text. Images have none by definition."""
    check_supported(path)
    ext = suffix(path)

    if ext == ".pdf":
        from extract_text import extract_text

        return extract_text(path)

    if ext == ".docx":
        import docx

        document = docx.Document(path)
        parts = [p.text for p in document.paragraphs if p.text.strip()]
        # Table cells also carry text worth searching, even though tables are
        # extracted separately with their structure intact.
        for table in document.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)

    if ext in TEXT_SUFFIXES:
        with open(path, encoding="utf-8", errors="replace") as handle:
            return handle.read()

    return ""  # a bare image


def load_tables(path: str) -> list[dict]:
    """Tables with their structure preserved, each with a caption for retrieval."""
    check_supported(path)
    ext = suffix(path)

    if ext == ".pdf":
        from extract_tables import extract_tables

        return extract_tables(path)

    if ext == ".docx":
        import docx

        from extract_tables import describe_from_headers, table_to_markdown

        document = docx.Document(path)
        tables = []
        for index, table in enumerate(document.tables, start=1):
            grid = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            if len(grid) < 2 or len(grid[0]) < 2:
                continue
            tables.append(
                {
                    # Word has no pages in the flow model; use table order instead.
                    "page": index,
                    "markdown": table_to_markdown(grid),
                    # Word tables rarely carry a machine-detectable caption, so
                    # describe them by their headers.
                    "caption": describe_from_headers(grid, index),
                }
            )
        return tables

    return []


def load_images(path: str, output_dir: str) -> list[dict]:
    """Figures to index. A bare image upload is itself the only figure."""
    check_supported(path)
    ext = suffix(path)

    if ext == ".pdf":
        from extract_images import extract_images

        return extract_images(path, output_dir)

    if ext in IMAGE_SUFFIXES:
        os.makedirs(output_dir, exist_ok=True)
        # Named like PDF-extracted figures so page parsing stays uniform downstream.
        target = os.path.join(output_dir, f"page1_img0{ext}")
        shutil.copyfile(path, target)
        return [{"page": 1, "path": target}]

    # python-docx does not expose inline images with position; skipped rather than
    # extracted out of order.
    return []
